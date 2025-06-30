from docx import Document
import os

def replace_text_in_header(header, old_text, new_text):
    replaced = False
    for paragraph in header.paragraphs:
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                replaced = True
    return replaced

# Folder witn initial files
input_folder = r'C:\Users\User\Desktop\Personal\CODING\my_proje\Magic substitution\files'

# Folder for saving new files
output_folder = r'C:\Users\User\Desktop\Personal\CODING\my_proje\Magic substitution\new_files'

old_text = 'Initial text'
new_text = 'TEST'

# going though all files in input_folder
for filename in os.listdir(input_folder):
    if filename.lower().endswith('.docx'):
        input_path = os.path.join(input_folder, filename)
        print(f'Processing file: {filename}')
        
        doc = Document(input_path)
        
        for p in doc.paragraphs:
            if '2024' in p.text:
                new_year = p.text.replace('2024','2025')
                for run in p.runs:
                    run.text = ''
                p.runs[0].text = new_year      
        
        overall_replaced = False
        for i, section in enumerate(doc.sections):
            header = section.header
            if replace_text_in_header(header, old_text, new_text):
                print(f"  Text is replaced in header of the section {i + 1}")
                overall_replaced = True
        
        if not overall_replaced:
            print("  No text to be substituted is found")
        
        # Naming of processed files adding 'new' before .docx
        name, ext = os.path.splitext(filename)
        output_filename = f"{name}_new{ext}"
        output_path = os.path.join(output_folder, output_filename)
                
        doc.save(output_path)
        # print("YEAR: ")
        # for i, paragraph in enumerate(doc.paragraphs[:20]):
        #     print(f"Параграф {i+1}: {paragraph.text}")
        
        
        print(f'Saved: {output_filename}\n')